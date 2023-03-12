import regex as re
from bs4 import BeautifulSoup
from html.parser import HTMLParser
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE

# ChatGPT 写一个python程序，需求是，输入一个html文档，把里面的一些字符，存成txt。需要提取的内容位于“题目内容”和“题目内容”之间，注意，内容不止一处。要去除html的各种标签。接下来还要做一些字符串处理。在"答对学生"，“答错学生”，“答题模式”，“题目类型“，”正确答案“和“未答学生”前面都要加一个换行符。数字前面的"："要去掉。
from bs4 import BeautifulSoup

# 读取HTML文档
with open('input.html', 'r', encoding='utf-8') as f:
    html = f.read()

# 使用BeautifulSoup库去除HTML标签
soup = BeautifulSoup(html, 'html.parser')
text = soup.get_text()

# ChatGPT 写一个python程序，需求是，输入一个txt文档，按照”未答学生“分割字符串，然后删除掉获取所有”答对学生“和”未答“之间的字符串

# 提取数字加句号和“答错学生”之间的字符
pattern = r'(\d+\.)\s*(.*?)未答学生'
matches = re.findall(pattern, text, re.S)
texts = [match[1].strip() for match in matches]

# 将提取的文本存储到列表中
text_list = list(texts)

def remove_text(lst):
    result = []
    for text in lst:
        start = text.find('答对学生')
        if start != -1:
            end = text.find('答错')
            if end != -1:
                result.append(text[:start] + text[end+2:])
            else:
                result.append(text[:start])
        else:
            start = text.find('答错')
            if start != -1:
                result.append(text[start+2:])
            else:
                result.append(text)
    return result
text_list = remove_text(text_list)
# 输出列表
print(text_list)

# ChatGPT: 写一个python程序，需求是，输入一个html文档，你需要提取里面所有的形式如同：2-4个汉字加上1-2个数字，的字符串。注意，这些字符串可能会重复，你要去除重复值。然后以每个字符串作为字典的键建立字典，字典的值设置为空字符串。

# 读取HTML文件
with open('input.html', 'r', encoding='utf-8') as f:
    html_doc = f.read()


# 使用BeautifulSoup解析HTML文档
soup = BeautifulSoup(html_doc, 'html.parser')

# 查找形式为“2-4个汉字加上1-2个数字”的字符串
pattern = r'[^\u4e00-\u9fa5\d]*(\p{Han}{2,4}\d{1,2})[^\u4e00-\u9fa5\d]*'
matches = []
for string in soup.stripped_strings:
    matches += re.findall(pattern, string, re.S)
unique_matches = list(set(matches))

# 创建字典
my_dict = {match: '' for match in unique_matches}


#ChatGPT: give me a python code. I will give you a dictionary and a list. The elements in the list are str. Then, for every key in the dictionary, you search for the key in every element of the list. If the key is in one element, you put the a part of element of the list into the value of the key. The part is the string before "答对学生" for every element in the list. If the key already has a value, append it.
# 输出字典
for key in my_dict:
    for element in text_list:
        if key in element:
            value = element.split("学生")[0].strip()
            if my_dict[key]:
                my_dict[key].append(value)
            else:
                my_dict[key] = [value]

print(my_dict)

# Export a dictionary to word file. The value of the key in the dictionary is a list. Therefore, you put every key in a line, every element in the list (i.e., dictionary values) in a line. Also, the key, when exporting, should be in red and bold.

doc = Document()

# Define red and bold styles for the key
key_format = doc.styles.add_style('Key Style', 1)
key_format.font.bold = True
key_format.font.color.rgb = RGBColor(255, 0, 0)

# Add each key and its values to the document
for key, values in my_dict.items():
    # Add the key to a new paragraph in red and bold
    p = doc.add_paragraph(style='Key Style')
    p.add_run(key)

    # Add each value in the list to a new paragraph
    for value in values:
        doc.add_paragraph(str(value))

# Save the document
doc.save("my_dictionary.docx")
